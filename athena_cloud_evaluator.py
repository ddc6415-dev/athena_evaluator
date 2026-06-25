import streamlit as st
import docx
import io
import requests
import re
from bs4 import BeautifulSoup
from youtube_transcript_api import YouTubeTranscriptApi

@st.cache_data
def write_report_to_word_bytes(report_text):
    """Generates and caches the .docx file for reliable downloading."""
    doc = docx.Document()
    doc.add_heading('Hexaxial Evaluation Report', 0)
    doc.add_paragraph(report_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.title("ATHENA-1: EVALUATION NODE")
    input_method = st.radio("Select input source", ["Local File (.docx)", "Web URL"])
    extracted_text = ""
    is_ready = False

    if input_method == "Web URL":
        target_url = st.text_input("Enter target URL")
        if st.button("Run Evaluation"):
            try:
                if "youtube.com" in target_url:
                    video_id = re.search(r"v=([a-zA-Z0-9_-]+)", target_url).group(1)
                    transcript_data = YouTubeTranscriptApi().get_transcript(video_id)
                    extracted_text = " ".join([t['text'] for t in transcript_data])
                else:
                    response = requests.get(target_url, headers={'User-Agent': 'Mozilla/5.0'})
                    extracted_text = BeautifulSoup(response.text, 'html.parser').get_text()
                is_ready = True
            except Exception as e:
                st.error(f"Input Processing Failed: {e}")
    else:
        uploaded_file = st.file_uploader("Upload Target Document", type=["docx"])
        if st.button("Run Evaluation"):
            if uploaded_file:
                try:
                    doc = docx.Document(uploaded_file)
                    extracted_text = "\n".join([p.text for p in doc.paragraphs])
                    is_ready = True
                except Exception as e:
                    st.error(f"File Processing Failed: {e}")

    if is_ready:
        st.success("Extraction complete.")
        with st.spinner("Processing evaluation via Gemini 2.5 Flash..."):
            try:
                api_key = st.secrets["GEMINI_API_KEY"]
                url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}"
                
                # ENHANCED PROMPT FOR STRICT FORMATTING COMPLIANCE
                payload = {
                    "contents": [{"parts": [{"text": (
                        f"You are the Author node of the Athena-1 system.\n\n"
                        f"Knowledge is information sufficiently accurate for repeated use.\n\n"
                        f"TASK 1 (Categorization): At the absolute beginning of your response, output ONLY the Operational Category on the first line. Example: [LEVEL: HYPOTHESIS / SPECULATIVE INFORMATION].\n\n"
                        f"TASK 2 (Graphic Display): Immediately following the category, output the exact header '-- HEXAXIAL METRIC DISTRIBUTION --'.\n\n"
                        f"TASK 3 (Metrics): Provide the score for all six axes using proportional Unicode blocks out of 10 (e.g., [██████░░░░]). Format exactly like the example below, with NO trailing text, explanations, or parentheses on the same line:\n"
                        f"Repeatability: [██████░░░░]\n"
                        f"Temporal Stability: [████████░░]\n"
                        f"Linguistic Precision: [██████████]\n"
                        f"Cultural Validation: [████░░░░░░]\n"
                        f"Technological Resolution: [████████░░]\n"
                        f"Sovereign Incentives: [██████░░░░]\n\n"
                        f"TASK 4 (Analysis): Below the graphical metrics, provide your textual executive summary and detailed historical analysis.\n\n"
                        f"TASK 5 (Advertising): Conclude with an [ADVERTISING_INTRUSION_LOG].\n\n"
                        f"--- PAYLOAD ---\n{extracted_text}"
                    )}]}]
                }
                
                response = requests.post(url, json=payload).json()
                
                if "candidates" in response:
                    result = response["candidates"][0]["content"]["parts"][0]["text"]
                    st.subheader("AI EVALUATION STREAM OUTPUT")
                    st.write(result)
                    
                    doc_buffer = write_report_to_word_bytes(result)
                    st.download_button(
                        label="Download Final Report", 
                        data=doc_buffer, 
                        file_name="Athena_Evaluation_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error(f"API Response Error: {response}")
            except Exception as e:
                st.error(f"Evaluation Failed: {e}")

if __name__ == "__main__":
    main()
